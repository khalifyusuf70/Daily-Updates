import os
import json
import sys
from datetime import datetime
from docx import Document
from io import BytesIO

# Fix: Better OpenAI import with version handling
try:
    from openai import OpenAI
except ImportError:
    import openai

# Railway will set this environment variable
PORT = int(os.environ.get("PORT", 8000))

# Initialize OpenAI with Railway environment variable
API_KEY = os.environ.get("OPENAI_API_KEY")
if not API_KEY:
    print("❌ OPENAI_API_KEY environment variable not set")
    print("Please set it in Railway dashboard: Variables -> OPENAI_API_KEY")
    sys.exit(1)

# Fix: Initialize client with proper error handling
try:
    # Try new version (>=1.0.0)
    client = OpenAI(api_key=API_KEY)
    print("✅ OpenAI client initialized (v1.0+)")
except TypeError:
    # Fallback to old version (<1.0.0)
    import openai
    openai.api_key = API_KEY
    print("✅ OpenAI client initialized (v0.x)")

def read_docx(file_path):
    """Extract text from a .docx file"""
    try:
        doc = Document(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        return "\n".join(paragraphs)
    except Exception as e:
        print(f"❌ Error reading {file_path}: {str(e)}")
        return ""

def write_docx(text, template_path, output_path):
    """Write tailored text to a new .docx preserving formatting"""
    try:
        doc = Document(template_path)
        
        # Split new text into paragraphs
        new_paragraphs = [p for p in text.split('\n') if p.strip()]
        
        # Replace paragraphs
        para_index = 0
        for paragraph in doc.paragraphs:
            if para_index >= len(new_paragraphs):
                break
            if paragraph.text.strip():
                if paragraph.runs:
                    paragraph.runs[0].text = new_paragraphs[para_index]
                    for run in paragraph.runs[1:]:
                        run.text = ""
                    para_index += 1
        
        # Add extra paragraphs if needed
        while para_index < len(new_paragraphs):
            doc.add_paragraph(new_paragraphs[para_index])
            para_index += 1
        
        doc.save(output_path)
        print(f"✅ Saved: {output_path}")
        return True
    except Exception as e:
        print(f"❌ Error writing {output_path}: {str(e)}")
        return False

def call_openai(prompt):
    """Unified OpenAI call for both versions"""
    try:
        # Try new version
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an expert CV tailoring assistant."},
                {"role": "user", "content": prompt}
            ],
            response_format={"type": "json_object"},
            temperature=0.3
        )
        return json.loads(response.choices[0].message.content)
    except AttributeError:
        # Fallback to old version
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an expert CV tailoring assistant."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3
        )
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        print(f"❌ OpenAI error: {str(e)}")
        return None

def tailor_cv(cv_text, job_description):
    """Send CV and job description to OpenAI for tailoring"""
    prompt = f"""
You are a professional CV tailoring assistant.

Input:
1. Master CV text
2. Job description text

Output in JSON format with one key:
- tailored_cv: The complete tailored CV text

Rules:
- NEVER add experience, skills, or qualifications that aren't in the master CV
- ONLY reword existing content to better match the job description
- Use keywords from the job description naturally
- Preserve all factual information

Master CV:
{cv_text}

Job Description:
{job_description}

Return JSON with the tailored CV:
"""
    
    result = call_openai(prompt)
    if result:
        return result.get("tailored_cv", "")
    return ""

def tailor_cover_letter(cover_text, job_description, cv_text):
    """Generate a tailored cover letter"""
    prompt = f"""
You are a professional cover letter writer.

Input:
1. Cover letter template
2. Job description
3. Master CV

Output in JSON format with one key:
- tailored_cover: The complete tailored cover letter

Rules:
- Use the template structure
- Personalize with specific examples from the CV
- Match keywords from the job description
- Keep it professional and concise

Cover Letter Template:
{cover_text}

Job Description:
{job_description}

Master CV:
{cv_text}

Return JSON with the tailored cover letter:
"""
    
    result = call_openai(prompt)
    if result:
        return result.get("tailored_cover", "")
    return ""

def process_application(cv_path, cover_path, job_description):
    """Process entire application tailoring"""
    # Read documents
    print("📄 Reading documents...")
    cv_text = read_docx(cv_path)
    cover_text = read_docx(cover_path)
    
    if not cv_text or not cover_text:
        print("❌ Failed to read documents")
        return None, None
    
    # Tailor CV
    print("🤖 Tailoring CV...")
    tailored_cv = tailor_cv(cv_text, job_description)
    
    if not tailored_cv:
        print("❌ Failed to tailor CV")
        return None, None
    
    # Tailor cover letter
    print("🤖 Tailoring Cover Letter...")
    tailored_cover = tailor_cover_letter(cover_text, job_description, cv_text)
    
    if not tailored_cover:
        print("❌ Failed to tailor cover letter")
        return None, None
    
    # Write documents
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    cv_output = f"Tailored_CV_{timestamp}.docx"
    cover_output = f"Tailored_Cover_{timestamp}.docx"
    
    cv_success = write_docx(tailored_cv, cv_path, cv_output)
    cover_success = write_docx(tailored_cover, cover_path, cover_output)
    
    if cv_success and cover_success:
        return cv_output, cover_output
    return None, None

def main():
    """Main entry point"""
    print("🚀 CV Tailor Service Starting...")
    print(f"📂 Working directory: {os.getcwd()}")
    print(f"📂 Files in directory: {os.listdir('.')}")
    
    # Check if files exist
    cv_file = "Master_CV.docx"
    cover_file = "Cover_Template.docx"
    
    if not os.path.exists(cv_file):
        print(f"❌ CV file not found: {cv_file}")
        print("Please upload your CV as 'Master_CV.docx'")
        return
    
    if not os.path.exists(cover_file):
        print(f"❌ Cover letter template not found: {cover_file}")
        print("Please upload your cover letter template as 'Cover_Template.docx'")
        return
    
    # Read job description from file
    job_desc_file = "job_description.txt"
    if os.path.exists(job_desc_file):
        with open(job_desc_file, 'r', encoding='utf-8') as f:
            job_description = f.read()
        print(f"📄 Loaded job description from {job_desc_file}")
    else:
        print("📝 Please paste the job description (type 'END' on new line to finish):")
        lines = []
        while True:
            try:
                line = input()
                if line.strip().upper() == "END":
                    break
                lines.append(line)
            except EOFError:
                break
        job_description = "\n".join(lines)
    
    if not job_description.strip():
        print("❌ No job description provided. Exiting.")
        return
    
    # Process application
    try:
        cv_output, cover_output = process_application(cv_file, cover_file, job_description)
        if cv_output and cover_output:
            print("✅ Tailored CV saved:", cv_output)
            print("✅ Tailored Cover Letter saved:", cover_output)
            print("🎉 Processing complete!")
        else:
            print("❌ Processing failed")
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
