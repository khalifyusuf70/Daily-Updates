import os
import json
import base64
import re
from datetime import datetime
from io import BytesIO
from flask import Flask, render_template, request, jsonify
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openai import OpenAI
import traceback

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "dev-secret-key-change-in-production")

# ---------------------------
# DEEPSEEK API CONFIGURATION
# ---------------------------
DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY")
if not DEEPSEEK_API_KEY:
    print("⚠️ DEEPSEEK_API_KEY not set")

client = OpenAI(
    api_key=DEEPSEEK_API_KEY,
    base_url="https://api.deepseek.com/v1"
)

COVER_PATH = "Cover_Template.docx"

# ---------------------------
# MASTER CV TEMPLATE
# ---------------------------
MASTER_CV_TEMPLATE = """
Chief of Staff (Feb 2023-To Date) | Jubaland State of Somalia
- Donor Engagement & Fundraising: Led resource mobilization efforts, engaging with international donors (USAID, EU, UN), diplomatic missions, and development partners to secure funding for state priorities and emergency response.
- Strategic Partnership Development: Established and maintained partnerships with international NGOs, UN agencies, and civil society organizations to enhance governance, humanitarian response, and service delivery.
- Strategic Leadership: Directed the Office of the President, coordinating activities across 15 government ministries and agencies to ensure alignment with the State Development Plan.
- Program Growth & Strategy: Led the development and implementation of the comprehensive State Development Plan, translating high-level goals into actionable programs and policies.
- Stakeholder Management: Represented the President at national and regional forums, managing complex relationships and ensuring aligned engagement across all partners.
- Crisis Management: Spearheaded emergency response efforts during political instability and humanitarian crises, coordinating relief activities and resource allocation.
- Established and maintained partnerships with civil society organizations and community leaders to enhance governance and service delivery.

Senior Advisor – Projects Planning & Grants Development (Aug 2021 – Jan 2023) | Ministry of Planning, Jubaland State, Somalia
- Grant Acquisition & Management: Led the development of the Ministry's Strategic Plan and managed a portfolio of donor-funded programs, ensuring effective implementation, financial accountability, and compliance.
- Proposal Development: Led humanitarian and development needs assessments, translating findings into priority program areas and successful funding proposals for international donors.
- Strategy & Policy: Provided strategic leadership, advising the Minister on planning, policy, and program development to align Ministry initiatives with national and regional priorities.
- External Relations: Established and maintained partnerships with international development organizations, UN agencies, and NGOs, fostering collaboration on civil society strengthening and peacebuilding initiatives.
- Analytical Reporting: Conducted political, economic, and social analysis to inform Ministry strategy, providing critical data-driven insights for programmatic decision-making.
- Capacity Building: Strengthened institutional capacity through training and technical support to ministry staff and partners.

Chief Operations Officer (Jul 2016 – Jul 2021) | KIMS MICROFINANCE, Somalia
- Organizational Leadership: Directed all operational functions, including budgeting, strategic planning, and resource allocation, to drive organizational development and growth.
- Business Development: Led business development initiatives, expanding the organization's reach into new regions and securing funding for microenterprise programming that supported vulnerable populations.
- Partnership Management: Established and maintained strategic partnerships with international development organizations to enhance program impact and sustainability.
- Community Engagement: Managed community relations, contributing to peacebuilding and social development through economic empowerment and locally-led programming.
- Stakeholder Relations: Represented the organization to government agencies, donors, and partners, ensuring effective communication and stakeholder management.
- Financial Oversight: Oversaw financial management, compliance, and reporting to ensure transparency and accountability.
- Team Leadership: Built and led high-performing teams, fostering a culture of excellence and continuous improvement.
"""

# ---------------------------
# EXACT JOB TITLES
# ---------------------------
JOB_TITLES = [
    "Chief of Staff (Feb 2023-To Date)",
    "Senior Advisor – Projects Planning & Grants Development (Aug 2021 – Jan 2023)",
    "Chief Operations Officer (Jul 2016 – Jul 2021)"
]

COMPANY_NAMES = {
    "Chief of Staff (Feb 2023-To Date)": "Jubaland State of Somalia",
    "Senior Advisor – Projects Planning & Grants Development (Aug 2021 – Jan 2023)": "Ministry of Planning, Jubaland State, Somalia",
    "Chief Operations Officer (Jul 2016 – Jul 2021)": "KIMS MICROFINANCE, Somalia"
}

# ---------------------------
# PROMPTS — plain executive CV bullets, no headings, no mapping phrases
# ---------------------------
SYSTEM_PROMPT = """
You are one of the world's leading Executive Resume Writers, Executive Recruiters,
ATS Optimization Specialists, and Career Strategists specializing in senior
international NGO, UN, Government, and Development Sector positions.

Your task is to tailor a candidate's CV to maximize alignment with a target job
description while maintaining complete honesty and factual accuracy.

CRITICAL RULES:
1. NEVER invent experience, achievements, employers, titles, or qualifications.
2. Rewrite only using truthful information from the candidate's CV.
3. Optimize for ATS by naturally incorporating exact keywords from the job description.
4. Write like a normal executive CV — do NOT use comparison language such as
   "directly aligned with", "parallel to", "mirrors", "equivalent to", "similar to",
   "relevant to", or "matching". Just state what the candidate did, in strong,
   confident, ATS-friendly language.
5. Do NOT create competency headings, themes, or categories (e.g. "Leadership",
   "Compliance", "Stakeholder Engagement", "Data Governance", "Risk Management").
   Every line must be a plain, standalone CV bullet.
6. Keep 6-8 bullet points per job.
7. Every bullet must begin with a strong action verb.

Example of the correct bullet style:
"Led strategic planning across government institutions to improve humanitarian programme delivery."
"Coordinated donor-funded programmes with UN agencies and development partners."
"Strengthened partnerships with civil society organisations to improve locally led programming."
"Conducted political and humanitarian analysis to inform executive decision-making."

The Professional Summary must:
- Be ONE paragraph of 5-6 concise sentences.
- Begin similarly to: "Highly efficient professional with over 10 years of senior
  leadership experience..."
- Match this style, tone and length (tailored to the specific job below):
  "Highly efficient professional with over 10 years of senior leadership experience
  managing complex portfolios across government and international development
  sectors. Proven ability to lead humanitarian response, civil society strengthening,
  and peacebuilding programs across the Horn of Africa. Expert in strategic
  leadership, programme growth, resource mobilization, and donor engagement. Strong
  track record in partnership development, locally led programming, and
  organizational transformation. Skilled in political analysis, contextual
  understanding, and managing multi-country operations."

Return ONLY valid JSON.
"""

def build_user_prompt(cv_text, job_description):
    return f"""
MASTER CV (candidate's actual experience):
{cv_text}

JOB DESCRIPTION (what the employer needs):
{job_description}

TASK:
Carefully analyse both documents. Tailor the CV specifically for this position while
maintaining complete factual accuracy.

RULES:
1. Keep 6-8 plain CV bullet points per job.
2. Do NOT use comparison/mapping phrases such as "directly aligned with", "parallel
   to", "mirrors", "equivalent to", "similar to", or "relevant to". Just describe
   what the candidate did, in strong executive CV language.
3. Do NOT create competency headings or categories inside a job (e.g. "Leadership:",
   "Compliance:", "Stakeholder Engagement:"). Every bullet must be a plain,
   standalone sentence — no bolded label before the colon.
4. Naturally weave in exact keywords from the job description.
5. Every bullet should begin with a strong action verb.
6. Never invent experience.

EXAMPLE of what a bullet should look like:
"Designed and delivered training for teams on SOPs for data collection, analysis, and management, tracking completion and reporting on gaps to senior leadership."
"Maintained and updated organizational templates and guidance materials to ensure alignment with the latest approved standards."
"Coordinated donor-funded programmes with UN agencies and development partners."

The "summary" field should be ONE executive paragraph of 5-6 concise sentences,
tailored specifically to this job, in the style described in the system prompt.

Return JSON:
{{
    "summary": "5-6 sentence executive summary tailored to this job",
    "experience": {{
        "Chief of Staff (Feb 2023-To Date)": [
            "plain CV bullet, no heading, no mapping phrase",
            "plain CV bullet, no heading, no mapping phrase",
            "plain CV bullet, no heading, no mapping phrase",
            "plain CV bullet, no heading, no mapping phrase",
            "plain CV bullet, no heading, no mapping phrase",
            "plain CV bullet, no heading, no mapping phrase",
            "plain CV bullet, no heading, no mapping phrase"
        ],
        "Senior Advisor – Projects Planning & Grants Development (Aug 2021 – Jan 2023)": [
            "plain CV bullet, no heading, no mapping phrase",
            "plain CV bullet, no heading, no mapping phrase",
            "plain CV bullet, no heading, no mapping phrase",
            "plain CV bullet, no heading, no mapping phrase",
            "plain CV bullet, no heading, no mapping phrase",
            "plain CV bullet, no heading, no mapping phrase"
        ],
        "Chief Operations Officer (Jul 2016 – Jul 2021)": [
            "plain CV bullet, no heading, no mapping phrase",
            "plain CV bullet, no heading, no mapping phrase",
            "plain CV bullet, no heading, no mapping phrase",
            "plain CV bullet, no heading, no mapping phrase",
            "plain CV bullet, no heading, no mapping phrase",
            "plain CV bullet, no heading, no mapping phrase"
        ]
    }}
}}
"""

def build_cover_letter_prompt(cv_text, job_description):
    return f"""
MASTER CV:
{cv_text}

JOB DESCRIPTION:
{job_description}

Write a 3-paragraph executive cover letter that:
1. Opens with enthusiasm and clearly states why the candidate is an excellent fit.
2. Connects the candidate's real experience to the job requirements in natural,
   confident prose — without stock comparison phrases like "directly aligned with",
   "parallel to", "mirrors", or "equivalent to".
3. Demonstrates understanding of the organization's mission.
4. Is persuasive, authentic, and sounds like a normal professional cover letter.
5. Avoids generic phrases.

Return ONLY the cover letter text.
"""

def call_ai(cv_text, job_description):
    """Call DeepSeek API with professional prompts"""
    try:
        print("📤 Sending request to DeepSeek API...")
        response = client.chat.completions.create(
            model="deepseek-v4-flash",
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": build_user_prompt(cv_text, job_description)},
            ],
            temperature=0.2,
            response_format={"type": "json_object"},
            timeout=45.0
        )
        
        raw_response = response.choices[0].message.content
        print(f"📥 Response received: {len(raw_response)} chars")
        
        result = json.loads(raw_response)
        
        if 'summary' not in result:
            result['summary'] = "Summary not provided"
        if 'experience' not in result:
            result['experience'] = {}
        
        return result
        
    except Exception as e:
        print(f"API error: {str(e)}")
        return {
            'summary': 'Error. Please try again.',
            'experience': {}
        }

def tailor_cover_letter(cv_text, job_description):
    """Generate tailored cover letter with explicit mapping"""
    try:
        print("✍️ Generating cover letter...")
        response = client.chat.completions.create(
            model="deepseek-v4-flash",
            messages=[
                {"role": "system", "content": "You are an expert cover letter writer for international organizations. Every sentence should connect the candidate's experience to the job requirements."},
                {"role": "user", "content": build_cover_letter_prompt(cv_text, job_description)}
            ],
            temperature=0.3,
            timeout=30.0
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"Cover letter error: {str(e)}")
        return "Dear Hiring Manager,\n\nI am writing to express my interest in this position. With over 10 years of experience in government and international development, I am confident in my ability to contribute to your organization.\n\nSincerely,\n[Your Name]"

def read_docx(file_path):
    """Extract text from .docx file"""
    try:
        doc = Document(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        return "\n".join(paragraphs)
    except Exception as e:
        print(f"Error reading {file_path}: {str(e)}")
        return ""

def create_tailored_docx(new_summary, new_experience):
    """Create tailored CV with skills section and 6-7 bullets per job"""
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(12)
    
    # Header
    doc.add_paragraph("+252611385559/+254715244144")
    doc.add_paragraph("cos.presidency@jubalandstate.so")
    doc.add_paragraph("khalifyusuf@agileanalytica.com")
    doc.add_paragraph("")
    
    # ========== SUMMARY ==========
    summary_header = doc.add_paragraph("# Summary")
    summary_header.runs[0].font.name = 'Calibri'
    summary_header.runs[0].font.size = Pt(12)
    summary_header.runs[0].bold = True
    
    if new_summary:
        summary_para = doc.add_paragraph(new_summary)
        summary_para.runs[0].font.name = 'Calibri'
        summary_para.runs[0].font.size = Pt(12)
    doc.add_paragraph("")
    
    # ========== SKILLS HIGHLIGHTS ==========
    skills_header = doc.add_paragraph("# Skill Highlights")
    skills_header.runs[0].font.name = 'Calibri'
    skills_header.runs[0].font.size = Pt(12)
    skills_header.runs[0].bold = True
    
    skills_left = [
        "• Resource Mobilization & Grants Management",
        "• Foundation Pipeline Management",
        "• Due Diligence & Risk Assessment",
        "• Impact Reporting & Communication",
        "• Political & Contextual Analysis"
    ]
    skills_right = [
        "• Team Leadership & Organizational Development",
        "• Risk Management & Compliance",
        "• Humanitarian Response Programming",
        "• Strategic Planning & Organizational Dev."
    ]
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.columns[0].width = Inches(3.5)
    table.columns[1].width = Inches(3.5)
    
    for row_idx in range(5):
        left_cell = table.cell(row_idx, 0)
        left_cell.text = skills_left[row_idx] if row_idx < len(skills_left) else ""
        for run in left_cell.paragraphs[0].runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
        
        right_cell = table.cell(row_idx, 1)
        right_cell.text = skills_right[row_idx] if row_idx < len(skills_right) else ""
        for run in right_cell.paragraphs[0].runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
    
    doc.add_paragraph("")
    
    # ========== EXPERIENCE ==========
    exp_header = doc.add_paragraph("# Experience")
    exp_header.runs[0].font.name = 'Calibri'
    exp_header.runs[0].font.size = Pt(12)
    exp_header.runs[0].bold = True
    
    for job_title in JOB_TITLES:
        # Job Title - Segoe UI 12 Bold
        job_para = doc.add_paragraph(job_title)
        job_para.runs[0].font.name = 'Segoe UI'
        job_para.runs[0].font.size = Pt(12)
        job_para.runs[0].bold = True
        
        # Company - Segoe UI 12 Italic
        company_para = doc.add_paragraph(COMPANY_NAMES.get(job_title, ""))
        company_para.runs[0].font.name = 'Segoe UI'
        company_para.runs[0].font.size = Pt(12)
        company_para.runs[0].italic = True
        
        # Bullets
        if job_title in new_experience and new_experience[job_title]:
            for bullet in new_experience[job_title]:
                bullet_para = doc.add_paragraph(f"• {bullet}")
                bullet_para.runs[0].font.name = 'Calibri'
                bullet_para.runs[0].font.size = Pt(12)
        else:
            # Fallback bullets
            fallback_bullets = {
                "Chief of Staff (Feb 2023-To Date)": [
                    "Led strategic coordination across 15 government ministries, directly aligned with coordinating research teams and cross-pillar collaboration.",
                    "Managed stakeholder relationships with international donors and NGOs, parallel to managing relationships with IRB and research leads.",
                    "Developed and implemented strategic plans, mirroring the development of SOPs and guidance materials.",
                    "Oversaw program monitoring and reporting, equivalent to supporting quarterly compliance monitoring.",
                    "Provided crisis management and emergency response coordination, similar to managing regulatory compliance and risk.",
                    "Facilitated resource mobilization and partnership building, relevant to organizing Research Community of Practice."
                ],
                "Senior Advisor – Projects Planning & Grants Development (Aug 2021 – Jan 2023)": [
                    "Led humanitarian needs assessments using research methodologies, directly aligned with designing training for researchers on data collection SOPs.",
                    "Managed donor-funded programs with strict compliance, parallel to maintaining research forms and templates for SOP compliance.",
                    "Conducted analysis and reporting for policy decisions, mirroring the analytical reporting required for IRB and research governance.",
                    "Established partnerships with UN agencies and NGOs, similar to organizing and coordinating Research Pillar-wide meetings.",
                    "Strengthened institutional capacity through training, directly matching the need to deliver training for researchers on SOPs.",
                    "Developed monitoring frameworks and reporting systems, equivalent to supporting quarterly compliance monitoring across the research portfolio."
                ],
                "Chief Operations Officer (Jul 2016 – Jul 2021)": [
                    "Directed operational functions including budgeting and planning, relevant to operating and maintaining AI-enabled research tools.",
                    "Led business development and market expansion, similar to piloting AI tools with research teams and developing guidance.",
                    "Established strategic partnerships, akin to coordinating cross-team knowledge exchange and the Research Community of Practice.",
                    "Managed stakeholder relations, relevant to organizing Research Pillar-wide meetings and managing IRB relationships.",
                    "Oversaw compliance and reporting, directly supporting SOP compliance monitoring.",
                    "Built and led teams, comparable to facilitating communities of practice and coordinating Data Champions."
                ]
            }
            for bullet in fallback_bullets.get(job_title, []):
                bullet_para = doc.add_paragraph(f"• {bullet}")
                bullet_para.runs[0].font.name = 'Calibri'
                bullet_para.runs[0].font.size = Pt(12)
        
        doc.add_paragraph("")
    
    # ========== EDUCATION ==========
    edu_header = doc.add_paragraph("# Education")
    edu_header.runs[0].font.name = 'Calibri'
    edu_header.runs[0].font.size = Pt(12)
    edu_header.runs[0].bold = True
    
    education = [
        "Micro Masters. Database management System",
        "University of Baltimore Maryland",
        "Degree Actuarial Science and Statistics",
        "JOMO KENYATTA UNIVERSITY",
        "Nairobi, Kenya"
    ]
    for edu in education:
        edu_para = doc.add_paragraph(edu)
        edu_para.runs[0].font.name = 'Calibri'
        edu_para.runs[0].font.size = Pt(12)
    
    doc.add_paragraph("")
    
    # ========== REFEREES ==========
    ref_header = doc.add_paragraph("# REFREES")
    ref_header.runs[0].font.name = 'Calibri'
    ref_header.runs[0].font.size = Pt(12)
    ref_header.runs[0].bold = True
    
    referees = [
        "Abshir olow", "Former Chief of Staff", "Jubaland State, Somalia",
        "+254722877178", "abshirmabdi@gmail.com", "",
        "Mr. Abdirahman Abdi", "Planning Minister, Jubaland",
        "+252612992848", "mopic@jubalandstate.so", "",
        "Omar Abdi Mohamed, HDmls, Msc.Int.",
        "Commodity Logistics Director-USAID", "omarabdi2@gmail.com"
    ]
    for ref in referees:
        if ref:
            ref_para = doc.add_paragraph(ref)
            ref_para.runs[0].font.name = 'Calibri'
            ref_para.runs[0].font.size = Pt(12)
        else:
            doc.add_paragraph()
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# ---------------------------
# MAIN PROCESSING
# ---------------------------
def process_application(job_description):
    """Process the entire application tailoring"""
    
    cv_text = MASTER_CV_TEMPLATE
    
    print(f"📄 CV text length: {len(cv_text)}")
    
    try:
        print("🤖 Tailoring CV with explicit mapping...")
        result = call_ai(cv_text, job_description)
        
        tailored_summary = result.get('summary', '')
        tailored_experience = result.get('experience', {})
        
        print(f"📝 Summary: {tailored_summary[:100] if tailored_summary else 'EMPTY'}...")
        print(f"📝 Experience keys: {list(tailored_experience.keys())}")
        
    except Exception as e:
        print(f"❌ AI error: {str(e)}")
        return None, None, f"AI error: {str(e)}", "Error", {}, "Error"
    
    try:
        print("✍️ Generating cover letter...")
        tailored_cover = tailor_cover_letter(cv_text, job_description)
        print(f"✅ Cover letter: {len(tailored_cover)} chars")
    except Exception as e:
        tailored_cover = "Dear Hiring Manager,\n\nI am writing to express my interest in this position. With over 10 years of experience in government and international development, I am confident in my ability to contribute to your organization.\n\nSincerely,\n[Your Name]"
        print(f"⚠️ Cover letter error: {str(e)}")
    
    try:
        print("📄 Creating DOCX...")
        cv_output = create_tailored_docx(tailored_summary, tailored_experience)
        
        cover_doc = Document()
        for line in tailored_cover.split('\n'):
            if line.strip():
                cover_doc.add_paragraph(line)
        cover_output = BytesIO()
        cover_doc.save(cover_output)
        cover_output.seek(0)
        
        print("✅ Documents generated")
        return cv_output, cover_output, "Success", tailored_summary, tailored_experience, tailored_cover
        
    except Exception as e:
        print(f"❌ Document error: {str(e)}")
        traceback.print_exc()
        return None, None, f"Document error: {str(e)}", "Error", {}, "Error"

# ---------------------------
# ROUTES
# ---------------------------
@app.route('/')
def index():
    try:
        return render_template('index.html')
    except Exception as e:
        return f"Error loading template: {str(e)}", 500

@app.route('/tailor', methods=['POST'])
def tailor():
    try:
        job_description = request.form.get('job_description', '').strip()
        if not job_description:
            return jsonify({'error': 'Please paste a job description'}), 400
        
        print(f"📝 Processing job: {len(job_description)} chars")
        cv_output, cover_output, message, tailored_summary, tailored_experience, tailored_cover = process_application(job_description)
        
        if cv_output and cover_output:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            cv_base64 = base64.b64encode(cv_output.getvalue()).decode('utf-8')
            cover_base64 = base64.b64encode(cover_output.getvalue()).decode('utf-8')
            
            return jsonify({
                'success': True,
                'message': 'Documents tailored successfully!',
                'cv_filename': f'Tailored_CV_{timestamp}.docx',
                'cover_filename': f'Tailored_Cover_{timestamp}.docx',
                'cv_data': cv_base64,
                'cover_data': cover_base64,
                'summary': tailored_summary,
                'experience': tailored_experience,
                'cover_letter': tailored_cover
            })
        else:
            return jsonify({'error': message}), 500
            
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/health')
def health():
    return jsonify({'status': 'healthy', 'timestamp': datetime.now().isoformat()})

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 8080))
    print(f"🚀 Starting CV Tailor on port {port}")
    print(f"📂 Files: {os.listdir('.')}")
    app.run(host='0.0.0.0', port=port, debug=False)
